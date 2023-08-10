#!/usr/bin/python3
# -*- coding:utf-8 -*-

remark = '''操作word文档'''

# 引入包
from docx import Document
from docx.shared import Inches
from Function.Common.function_common import *
class Word:

    def __init__(self,doc):
        self.doc = doc

    # level可以为0-9，其中0是带有下划线样式的标题
    def add_heading(self,content,level):
        self.doc.add_heading(content,level)

    # 保存文档
    def save_word(self,filename):
        self.doc.save(filename)





# '''
# win32com在word中插入图片(https://www.zhihu.com/column/c_1365074188886163456)
#
# import win32com.client as win32
# from win32com.client import constants
# # 1）打开word应用程序
# doc_app = win32.gencache.EnsureDispatch('Word.Application')
# doc_app.Visible =1
# # 2）添加一个新得word文档
# doc = doc_app.Documents.Add()
# # 3）添加新的段落
# parag = doc.Paragraphs.Add()
# parag_range = parag.Range
# parag_range.Text ='新插入的图片'#随便插入一段文字
# #指定文件的完整路径
# picture_full_path = r'C:\Users\XXXX\Pictures\cat.jpg'
# #在当前的段落中插入图片
# parag_range.InlineShapes.AddPicture(picture_full_path)
# '''

doc = Document()
doc.add_heading('一级标题',level=1)
p = doc.add_paragraph('', style='List Bullet')
p.add_run('aaaaa')
doc.add_picture(r'E:\TDdownload\名不虚传.jpg', width=Inches(5.0))
p = doc.add_paragraph('', style='List Bullet')
doc.save(r'E:\TDdownload\aaa.docx')
import webbrowser
import win32api
import win32con
webbrowser.open(r'E:\TDdownload\aaa.docx')
# f=open(r'E:\TDdownload\aaa.docx')
# f.seek(0,2)
time.sleep(4)
# handle = win32gui.FindWindow("aaa.docx", None)
#ctrl+End移动到word末尾
win32api.keybd_event(17, 0, 0, 0)
win32api.keybd_event(35, 0, 0, 0)
win32api.keybd_event(35, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(17, 0, win32con.KEYEVENTF_KEYUP, 0)
#输入回车
win32api.keybd_event(13, 0, 0, 0)
win32api.keybd_event(13, 0, win32con.KEYEVENTF_KEYUP, 0)
#crtl+N打开拆入对象窗口
win32api.keybd_event(18, 0, 0, 0)
win32api.keybd_event(78, 0, 0, 0)
win32api.keybd_event(78, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(74, 0, 0, 0)
win32api.keybd_event(74, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(74, 0, 0, 0)
win32api.keybd_event(74, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, 0, 0)
win32api.keybd_event(70, 0, 0, 0)
win32api.keybd_event(70, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, 0, 0)
win32api.keybd_event(66, 0, 0, 0)
win32api.keybd_event(66, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, win32con.KEYEVENTF_KEYUP, 0)
WinUpLoadFile().winUpLoadFile(r'E:\TDdownload\Q-Dir_Installer_x64.zip','浏览')
win32api.keybd_event(18, 0, 0, 0)
win32api.keybd_event(65, 0, 0, 0)
win32api.keybd_event(65, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(18, 0, win32con.KEYEVENTF_KEYUP, 0)
win32api.keybd_event(13, 0, 0, 0)
win32api.keybd_event(13, 0, win32con.KEYEVENTF_KEYUP, 0)
# import zipfile,shutil
# from random import randint
# filename=[r'E:\TDdownload\aaa.docx',r'D:\TDdownload\Document\TestFile\word文档1.docx',r'D:\TDdownload\Document\TestFile\word文档2.docx']
#
# #filename为相关文件列表，filename[0]就是需要处理的主word文件，filename[1]-[12]是12个需要插入的附件
# import os,zipfile,shutil       #需要用到的包
# azip = zipfile.ZipFile(r'E:\TDdownload\aaa.docx')
# #以压缩格式打开word文件
# tempdir=''
# while True:
#     tempdir= str(randint(11111111,99999999))  #生行8位临时文件夹名
#     if not os.path.exists(tempdir):
#         break
# os.mkdir(tempdir)                  #创建临时目录
# os.chdir(tempdir)                   #转到临时目录
# azip.extractall()                     #解压word文件到临时文件夹
# azip.close()                           #关闭word文档，否则后面重新压缩会报错
# #把正确文件拷贝覆盖模版文件的空附件
# try:
#     shutil.copy(filename[1],'word\\embeddings\\Microsoft_Word_Document.docx')
#     shutil.copy(filename[2],'word\\embeddings\\Microsoft_Word_Document1.docx')
#     azip = zipfile.ZipFile(filename[0], 'w')    #以压缩格式新建word文档
#     for i in os.walk('.'):                             #使用os.walk遍历整个目录及子目录，保证原有的目录结构不变
#         for j in i[2]:
#             azip.write(os.path.join(i[0],j), compress_type=zipfile.ZIP_DEFLATED)     #将文件逐个打包到word文档中，压缩格式指定为ZIP_DEFLATED
#     azip.close()                                       #关闭文件
#     os.chdir('..')
#     shutil.rmtree(tempdir,ignore_errors=True)    #删除临时文件夹
# except:
#     pass

doc = Document(r'E:\TDdownload\aaa.docx')
doc.paragraphs[0].text










































if __name__ == "__main__":
    # doc = Document()
    # Word(doc).add_heading('一级标题',level=1)
    # Word(doc).save_word(r'E:\TDdownload\aaa.docx')
    pass